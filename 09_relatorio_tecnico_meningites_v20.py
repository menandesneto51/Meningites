# -*- coding: utf-8 -*-
"""
09_relatorio_tecnico_meningites_v20.py
Relatório técnico V20 usando novos módulos: OR por classificação, distância, laboratório real e score de qualidade.
"""

import numpy as np
import pandas as pd
from pathlib import Path
from meningites_v17_common import *

def read(path):
    p = Path(path)
    if not p.exists():
        return pd.DataFrame()
    return pd.read_csv(p, encoding="utf-8-sig", low_memory=False)

def safe_head(df, n=10):
    return df.head(n) if not df.empty else df

def main():
    base = read(OUT / "base_unica_meningites_v17.csv")
    kpi = read(OUT / "kpis_semanais_v17.csv")
    or_class = read(OUT / "odds_classificacao_desfechos_v20.csv")
    lab_kpi = read(OUT / "laboratorio_kpis_v20.csv")
    lab_class = read(OUT / "indicadores_laboratoriais_classificacao_v20.csv")
    dist = read(OUT / "geoespacial_laboratorio_distancia_v20.csv")
    corr_dist = read(OUT / "correlacao_distancia_laboratorio_v20.csv")
    qres = read(OUT / "qualidade_score_resumo_v20.csv")
    qscore = read(OUT / "qualidade_score_v20.csv")
    com = read(OUT / "associacoes_comorbidades_quiquadrado_v18.csv")
    alerts = read(OUT / "top_alertas_surtos_v17.csv")
    fc = read(OUT / "forecasting_resumo_v17.csv")
    moran = read(OUT / "moran_global_v17.csv")

    lines = []
    lines.append("# RELATÓRIO TÉCNICO – MENINGITES CIEVS/MT – V20")
    lines.append("")
    lines.append("## 1. Cenário atual")
    if not base.empty:
        casos = len(base)
        confirmados = int(pd.to_numeric(base.get("confirmado_v17", 0), errors="coerce").sum())
        hosp = int(pd.to_numeric(base.get("hospitalizacao_v17", 0), errors="coerce").sum())
        obitos = int(pd.to_numeric(base.get("obito_meningite_v17", 0), errors="coerce").sum())
        lines.append(f"Foram analisados {fmt_num(casos)} registros, com {fmt_num(confirmados)} confirmados, {fmt_num(hosp)} hospitalizações e {fmt_num(obitos)} óbitos por meningite.")
        if "classificacao_agrupada_v17" in base.columns:
            top = base["classificacao_agrupada_v17"].value_counts().head(10)
            lines.append("Distribuição por classificação agrupada: " + "; ".join([f"{k}: {fmt_num(v)}" for k, v in top.items()]) + ".")
    else:
        lines.append("Base V20 indisponível.")

    lines.append("")
    lines.append("## 2. KPIs semanais")
    if not kpi.empty:
        for _, r in kpi.iterrows():
            lines.append(f"- {r['indicador_rotulo']}: semana atual {fmt_num(r['valor_atual_fechado'])}; semana anterior {fmt_num(r['valor_semana_anterior'])}; variação {fmt_num(r['variacao_percentual'])}%; semáforo {r['semaforo']}.")
    else:
        lines.append("KPIs semanais indisponíveis.")

    lines.append("")
    lines.append("## 3. Odds Ratio por classificação agrupada")
    if not or_class.empty:
        or_class["p_value"] = pd.to_numeric(or_class["p_value"], errors="coerce")
        show = or_class.sort_values(["desfecho", "p_value"]).head(20)
        for _, r in show.iterrows():
            lines.append(f"- {r['classificacao_agrupada']} | {r['desfecho']}: OR={fmt_num(r['or'],2)} (IC95% {fmt_num(r['ic95_inferior'],2)}–{fmt_num(r['ic95_superior'],2)}), p={fmt_num(r['p_value'],4)}. {r['interpretacao_estatistica']} {r['relevancia_pratica']}")
    else:
        lines.append("OR por classificação agrupada indisponível.")

    lines.append("")
    lines.append("## 4. Geoespacial, distância de Cuiabá e uso do laboratório")
    if not moran.empty:
        for _, r in moran.iterrows():
            lines.append(f"- Moran: I={fmt_num(r.get('moran_i', np.nan),3)}; p={fmt_num(r.get('p_value', np.nan),4)}. {r.get('interpretacao','')}")
    if not corr_dist.empty:
        for _, r in corr_dist.iterrows():
            lines.append(f"- Distância x {r['variavel']}: Spearman={fmt_num(r['spearman_distancia'],3)}; n={fmt_num(r['n'],0)}.")
    if not dist.empty:
        top_dist = dist.sort_values("distancia_cuiaba_km", ascending=False).head(5) if "distancia_cuiaba_km" in dist.columns else dist.head(5)
        lines.append("Municípios mais distantes com avaliação laboratorial: " + "; ".join([f"{r.get('municipio_v17','NA')} ({fmt_num(r.get('distancia_cuiaba_km', np.nan))} km; uso lab {fmt_num(r.get('taxa_uso_laboratorio_pct', np.nan))}%)" for _, r in top_dist.iterrows()]) + ".")

    lines.append("")
    lines.append("## 5. Laboratório")
    if not lab_kpi.empty:
        r = lab_kpi.iloc[0]
        lines.append(f"Total de notificações: {fmt_num(r.get('total_notificacoes',np.nan))}; confirmados: {fmt_num(r.get('total_confirmados',np.nan))}; resultados laboratoriais concludentes: {fmt_num(r.get('total_com_resultado_laboratorial_concludente',np.nan))}; taxa de positividade real: {fmt_num(r.get('taxa_positividade_real_pct',np.nan))}%.")
    if not lab_class.empty:
        for _, r in lab_class.sort_values("taxa_positividade_real_pct", ascending=False).head(10).iterrows():
            lines.append(f"- {r['classificacao_agrupada_v17']}: positividade real {fmt_num(r['taxa_positividade_real_pct'])}% e cobertura laboratorial {fmt_num(r['cobertura_laboratorial_pct'])}%.")

    lines.append("")
    lines.append("## 6. Comorbidades")
    if not com.empty:
        com["p_value"] = pd.to_numeric(com["p_value"], errors="coerce")
        for _, r in com.sort_values("p_value").head(10).iterrows():
            lines.append(f"- {r['variavel']} vs {r['desfecho']}: p={fmt_num(r['p_value'],4)}; Cramér's V={fmt_num(r['cramers_v'],3)}. {r['interpretacao_estatistica']}")
    else:
        lines.append("Análise de comorbidades indisponível.")

    lines.append("")
    lines.append("## 7. Projeções")
    if not fc.empty:
        for _, r in fc.head(20).iterrows():
            lines.append(f"- {r['desfecho']} em {fmt_num(r['horizonte_dias'],0)} dias: {fmt_num(r['valor_previsto'])} (IC95% {fmt_num(r['ic95_inferior'])}–{fmt_num(r['ic95_superior'])}).")

    lines.append("")
    lines.append("## 8. Qualidade do banco")
    if not qres.empty:
        r = qres.iloc[0]
        lines.append(f"Pontuação total: {fmt_num(r['pontuacao_total'],0)}/20. Classificação: {r['qualidade_banco']}.")
    if not qscore.empty:
        for _, r in qscore.iterrows():
            lines.append(f"- {r['criterio']}: {fmt_num(r['pontuacao'],0)} ponto(s). {r['interpretacao']}")

    lines.append("")
    lines.append("## 9. Alertas de surtos")
    if not alerts.empty:
        for _, r in alerts.head(10).iterrows():
            lines.append(f"- {r.get('municipio_v17','NA')} | {r.get('classificacao_agrupada_v17','NA')} | {r.get('classe_alerta','NA')} | {r.get('motivos','')}")
    else:
        lines.append("Nenhum alerta acima de rotina disponível no arquivo de saída.")

    lines.append("")
    lines.append("## 10. Recomendações estratégicas")
    lines.append("- Usar os KPIs da aba 01 para leitura do período selecionado, mantendo a comparação da última semana como sinal operacional.")
    lines.append("- Priorizar investigações conforme classificação agrupada, OR, surtos e qualidade do dado.")
    lines.append("- Revisar acesso laboratorial de municípios mais distantes de Cuiabá quando houver baixa cobertura laboratorial.")
    lines.append("- Monitorar positividade real por classificação agrupada e critério de confirmação.")
    lines.append("- Usar a pontuação de qualidade 0–20 como síntese para orientar limpeza, busca ativa e qualificação do banco.")

    REL.mkdir(exist_ok=True)
    md = REL / "RELATORIO_TECNICO_MENINGITES_CIEVS_MT_V20.md"
    docx = REL / "RELATORIO_TECNICO_MENINGITES_CIEVS_MT_V20.docx"
    md.write_text("\n".join(lines), encoding="utf-8")

    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(11)
        for line in lines:
            if line.startswith("# "):
                doc.add_heading(line.replace("# ", ""), level=1)
            elif line.startswith("## "):
                doc.add_heading(line.replace("## ", ""), level=2)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:])
            else:
                doc.add_paragraph(line)
        doc.save(docx)
        print("[OK]", docx)
    except Exception as e:
        print("[AVISO] DOCX não gerado:", e)
    print("[OK]", md)

if __name__ == "__main__":
    main()
